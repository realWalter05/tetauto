from docx import Document
from docx.shared import Pt, Mm
from datetime import datetime
import os


class Doc:


    def __init__(self):
        self.document = Document()
        self.setup()


    def setup(self):
        section = self.document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)


    def save_document(self):
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        self.document.save(os.path.join(desktop, "SOUHLAS ZAMĚSTNANCE SE ZASÍLÁNÍM ELEKTRONICKÉ VÝPLATNÍ PÁSKY.docx"))


    def add_page(self, person):
        self.document.add_paragraph("\n\n\n")

        title = self.document.add_paragraph()
        title_text = "SOUHLAS ZAMĚSTNANCE SE ZASÍLÁNÍM ELEKTRONICKÉ VÝPLATNÍ PÁSKY"
        t = title.add_run(title_text)
        t.bold = True
        t.underline = True
        title.alignment = 1
        t.font.size = Pt(13)

        self.document.add_paragraph("\n")
        self.document.add_paragraph(f"Jméno a příjmení:").add_run(f" {person.name}").bold = True
        self.document.add_paragraph(f"Osobní číslo:").add_run(f"          {str(person.id)}").bold = True
        self.document.add_paragraph(f"Adresa:                    ").add_run(f"{person.street}").bold = True

        self.document.add_paragraph("\n\n\n")
        self.document.add_paragraph("Souhlasím se zasíláním elektronické výplatní pásky (mzdového lístku) na e-mailovou adresu:")


        t = self.document.add_paragraph()
        t.alignment = 1
        title = t.add_run(person.email)
        title.bold = True
        title.font.size = Pt(12)

        self.document.add_paragraph("\n\n\n")

        self.document.add_paragraph("Případné změny e-mailové adresy neprodleně oznámím svému zaměstnavateli:")
        self.document.add_paragraph("      Domov pro seniory Světlo, Drhovle 44 – Zámek, 397 01 Písek.")

        self.document.add_paragraph("\n\n\n\n")

        date_today = datetime.now().strftime("%d.%m.%Y")
        self.document.add_paragraph(f"V Drhovli {date_today}")
        self.document.add_paragraph("")

        self.document.add_paragraph("                                                                                                                    podpis zaměstnance")
        self.document.add_page_break()

